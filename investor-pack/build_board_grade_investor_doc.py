from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUTPUT_FILE = Path(__file__).resolve().parent / "CLISONIX_Board_Grade_Investor_Memo_EN_SQ_DE.docx"


def title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_3y_table(doc: Document, labels: list[str]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, value in enumerate(labels):
        header[idx].text = value

    rows = [
        ["€424,500", "€1,862,000", "€5,230,000"],
        ["€701,004", "€1,089,872", "€1,821,180"],
        ["-€276,504", "€772,128", "€3,408,820"],
        ["-65.1%", "41.5%", "65.2%"],
    ]

    for row_values in rows:
        row = table.add_row().cells
        row[0].text = labels[0]
        row[1].text = row_values[0]
        row[2].text = row_values[1]
        row[3].text = row_values[2]
        labels[0] = {
            "Metric": "Total Expenses",
            "Treguesi": "Shpenzimet Totale",
            "Kennzahl": "Gesamtkosten",
        }.get(labels[0], "Total Expenses") if "Revenue" in row[0].text or "Ardhurat" in row[0].text or "Umsatz" in row[0].text else labels[0]


def add_financial_table(doc: Document, lang: str) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    if lang == "EN":
        headers = ["Metric", "2026", "2027", "2028"]
        data = [
            ["Total Revenue", "€424,500", "€1,862,000", "€5,230,000"],
            ["Total Expenses", "€701,004", "€1,089,872", "€1,821,180"],
            ["EBITDA", "-€276,504", "€772,128", "€3,408,820"],
            ["EBITDA Margin", "-65.1%", "41.5%", "65.2%"],
        ]
    elif lang == "SQ":
        headers = ["Treguesi", "2026", "2027", "2028"]
        data = [
            ["Të Ardhurat Totale", "€424,500", "€1,862,000", "€5,230,000"],
            ["Shpenzimet Totale", "€701,004", "€1,089,872", "€1,821,180"],
            ["EBITDA", "-€276,504", "€772,128", "€3,408,820"],
            ["Marzhi EBITDA", "-65.1%", "41.5%", "65.2%"],
        ]
    else:
        headers = ["Kennzahl", "2026", "2027", "2028"]
        data = [
            ["Gesamtumsatz", "€424,500", "€1,862,000", "€5,230,000"],
            ["Gesamtkosten", "€701,004", "€1,089,872", "€1,821,180"],
            ["EBITDA", "-€276,504", "€772,128", "€3,408,820"],
            ["EBITDA-Marge", "-65.1%", "41.5%", "65.2%"],
        ]

    hdr = table.rows[0].cells
    for idx, val in enumerate(headers):
        hdr[idx].text = val

    for row_values in data:
        row = table.add_row().cells
        for idx, val in enumerate(row_values):
            row[idx].text = val


def english_section(doc: Document) -> None:
    h2(doc, "English | Board-Grade Memorandum")
    doc.add_paragraph(
        "Clisonix is live at clisox.com and used by users across multiple countries. This memorandum is prepared for institutional and venture capital diligence in Germany and the UK."
    )
    h2(doc, "Investment Case")
    bullets(
        doc,
        [
            "Validated live platform with real user activity and production uptime.",
            "Differentiated stack: ASI Trinity + Ocean Core orchestration + Excel Core reporting.",
            "Revenue model combining self-serve SaaS and higher-value research/enterprise plans.",
            "Disciplined infrastructure economics with Hetzner GPU baseline and controlled backup costs.",
        ],
    )
    h2(doc, "3-Year Financial Projection")
    add_financial_table(doc, "EN")
    h2(doc, "Legal and Financial Note")
    doc.add_paragraph(
        "This document contains management projections and forward-looking statements prepared for strategic evaluation. It does not constitute a public offer, investment advice, or a binding commitment."
    )


def albanian_section(doc: Document) -> None:
    doc.add_page_break()
    h2(doc, "Shqip | Memorandum Board-Grade")
    doc.add_paragraph(
        "Clisonix është live në clisox.com dhe përdoret nga përdorues në disa vende të botës. Ky memorandum është përgatitur për due diligence nga investitorë institucionalë dhe VC në Gjermani dhe UK."
    )
    h2(doc, "Teza e Investimit")
    bullets(
        doc,
        [
            "Platformë reale në prodhim me aktivitet përdoruesish.",
            "Stack i diferencuar: ASI Trinity + orkestrimi Ocean Core + raportimi Excel Core.",
            "Model i kombinuar monetizimi: SaaS self-serve + plane Research/Enterprise.",
            "Efikasitet i kostos së infrastrukturës me bazë Hetzner GPU dhe kosto backup të kontrolluar.",
        ],
    )
    h2(doc, "Projeksioni Financiar 3-Vjeçar")
    add_financial_table(doc, "SQ")
    h2(doc, "Shënim Ligjor dhe Financiar")
    doc.add_paragraph(
        "Ky dokument përmban projeksione menaxheriale dhe deklarata orientuese për të ardhmen, të përdorshme vetëm për vlerësim strategjik. Nuk përbën ofertë publike, këshillim investimi apo angazhim detyrues."
    )


def german_section(doc: Document) -> None:
    doc.add_page_break()
    h2(doc, "Deutsch | Board-Grade Memorandum")
    doc.add_paragraph(
        "Clisonix ist live auf clisox.com und wird bereits von Nutzern in mehreren Ländern verwendet. Dieses Memorandum wurde für die Due-Diligence-Prüfung von institutionellen Investoren und VCs in Deutschland und UK erstellt."
    )
    h2(doc, "Investment-These")
    bullets(
        doc,
        [
            "Validierte Live-Plattform mit realer Nutzeraktivität.",
            "Differenzierter Stack: ASI Trinity + Ocean Core Orchestrierung + Excel Core Reporting.",
            "Monetarisierung durch Self-Serve-SaaS sowie Research- und Enterprise-Verträge.",
            "Disziplinierte Infrastrukturkosten mit Hetzner-GPU-Basis und kontrollierten Backup-Kosten.",
        ],
    )
    h2(doc, "3-Jahres-Finanzprojektion")
    add_financial_table(doc, "DE")
    h2(doc, "Rechtlicher und Finanzieller Hinweis")
    doc.add_paragraph(
        "Dieses Dokument enthält Management-Projektionen und zukunftsgerichtete Aussagen für strategische Bewertungszwecke. Es stellt weder ein öffentliches Angebot noch eine Anlageberatung oder eine bindende Zusage dar."
    )


def main() -> None:
    doc = Document()
    title(doc, "CLISONIX — Board-Grade Investor Memorandum")
    subtitle(doc, "Formal Version for VC Review (Germany/UK) | English • Shqip • Deutsch")
    doc.add_paragraph("Date: March 2026")
    doc.add_paragraph("Reference files: investor-pack/CLISONIX_Financial_Model_Investor_Ready.xlsx")

    english_section(doc)
    albanian_section(doc)
    german_section(doc)

    doc.save(OUTPUT_FILE)
    print(f"Generated: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
