from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUTPUT_FILE = Path(__file__).resolve().parent / "CLISONIX_Investor_Document_EN_SQ_DE.docx"


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)


def add_h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_financial_table(doc: Document, lang: str) -> None:
    if lang == "EN":
        headers = ["Metric", "2026", "2027", "2028"]
        rows = [
            ["Total Revenue", "€424,500", "€1,862,000", "€5,230,000"],
            ["Total Expenses", "€701,004", "€1,089,872", "€1,821,180"],
            ["EBITDA", "-€276,504", "€772,128", "€3,408,820"],
            ["EBITDA Margin", "-65.1%", "41.5%", "65.2%"],
        ]
    elif lang == "SQ":
        headers = ["Treguesi", "2026", "2027", "2028"]
        rows = [
            ["Të Ardhurat Totale", "€424,500", "€1,862,000", "€5,230,000"],
            ["Shpenzimet Totale", "€701,004", "€1,089,872", "€1,821,180"],
            ["EBITDA", "-€276,504", "€772,128", "€3,408,820"],
            ["Marzhi EBITDA", "-65.1%", "41.5%", "65.2%"],
        ]
    else:
        headers = ["Kennzahl", "2026", "2027", "2028"]
        rows = [
            ["Gesamtumsatz", "€424,500", "€1,862,000", "€5,230,000"],
            ["Gesamtkosten", "€701,004", "€1,089,872", "€1,821,180"],
            ["EBITDA", "-€276,504", "€772,128", "€3,408,820"],
            ["EBITDA-Marge", "-65.1%", "41.5%", "65.2%"],
        ]

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, head in enumerate(headers):
        hdr[i].text = head
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = value


def build_english_section(doc: Document) -> None:
    doc.add_page_break()
    add_h2(doc, "English Version")
    doc.add_paragraph(
        "Clisonix is a neural intelligence platform with 17+ active modules across AI reasoning, chat, research, neuroscience, health, environment, data, and developer tools."
    )
    add_h2(doc, "Investment Highlights")
    add_bullets(
        doc,
        [
            "Live platform with 99.97% uptime and multi-module product readiness.",
            "Clear monetization model: Free, Pro (€29), Research (€99), Enterprise (€499+), API/Developer (€299).",
            "Efficient infrastructure baseline: Hetzner RTX 5090 (€509/month) + storage/backups (€258/month).",
            "Scalable unit economics with improving CAC and LTV/CAC trajectory.",
        ],
    )
    add_h2(doc, "3-Year Financial Projection")
    add_financial_table(doc, "EN")
    add_h2(doc, "Seed Round")
    doc.add_paragraph("Capital sought: €600,000")
    add_bullets(
        doc,
        [
            "Product Development: 40%",
            "Marketing Launch: 30%",
            "Infrastructure: 10%",
            "Team Expansion: 15%",
            "Operations & Buffer: 5%",
        ],
    )


def build_albanian_section(doc: Document) -> None:
    doc.add_page_break()
    add_h2(doc, "Versioni Shqip")
    doc.add_paragraph(
        "Clisonix është platformë neural intelligence me 17+ module aktive në AI reasoning, chat, research, neuroscience, health, environment, data dhe developer tools."
    )
    add_h2(doc, "Pikat Kryesore për Investim")
    add_bullets(
        doc,
        [
            "Platformë live me 99.97% uptime dhe module production-ready.",
            "Model i qartë monetizimi: Free, Pro (€29), Research (€99), Enterprise (€499+), API/Developer (€299).",
            "Infrastrukturë efikase: Hetzner RTX 5090 (€509/muaj) + storage/backups (€258/muaj).",
            "Unit economics në përmirësim me trend pozitiv CAC dhe LTV/CAC.",
        ],
    )
    add_h2(doc, "Projeksioni Financiar 3-Vjeçar")
    add_financial_table(doc, "SQ")
    add_h2(doc, "Seed Round")
    doc.add_paragraph("Kapitali i kërkuar: €600,000")
    add_bullets(
        doc,
        [
            "Product Development: 40%",
            "Marketing Launch: 30%",
            "Infrastructure: 10%",
            "Team Expansion: 15%",
            "Operations & Buffer: 5%",
        ],
    )


def build_german_section(doc: Document) -> None:
    doc.add_page_break()
    add_h2(doc, "Deutsche Version")
    doc.add_paragraph(
        "Clisonix ist eine Neural-Intelligence-Plattform mit mehr als 17 aktiven Modulen in den Bereichen AI Reasoning, Chat, Research, Neuroscience, Health, Environment, Data und Developer Tools."
    )
    add_h2(doc, "Investment Highlights")
    add_bullets(
        doc,
        [
            "Live-Plattform mit 99,97% Uptime und produktionsreifen Modulen.",
            "Klares Monetarisierungsmodell: Free, Pro (€29), Research (€99), Enterprise (€499+), API/Developer (€299).",
            "Effiziente Infrastruktur: Hetzner RTX 5090 (€509/Monat) + Storage/Backups (€258/Monat).",
            "Skalierbare Unit Economics mit verbessertem CAC und LTV/CAC-Verhältnis.",
        ],
    )
    add_h2(doc, "3-Jahres-Finanzprojektion")
    add_financial_table(doc, "DE")
    add_h2(doc, "Seed-Runde")
    doc.add_paragraph("Gesuchtes Kapital: €600.000")
    add_bullets(
        doc,
        [
            "Produktentwicklung: 40%",
            "Marketing Launch: 30%",
            "Infrastruktur: 10%",
            "Teamerweiterung: 15%",
            "Operations & Buffer: 5%",
        ],
    )


def main() -> None:
    doc = Document()

    add_title(doc, "CLISONIX — Investor Document")
    add_subtitle(doc, "Professional Trilingual Version (English | Shqip | Deutsch)")

    doc.add_paragraph("Prepared for investor outreach and due diligence discussions.")
    doc.add_paragraph("Date: March 2026")

    build_english_section(doc)
    build_albanian_section(doc)
    build_german_section(doc)

    doc.save(OUTPUT_FILE)
    print(f"Generated: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
